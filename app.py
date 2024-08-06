from flask import Flask, render_template, jsonify, flash, redirect, url_for, session, request, logging
from flask import send_file, abort
from flask_mail import Mail, Message
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy.orm import aliased
from sqlalchemy import func
from sqlalchemy import not_
from wtforms import Form, StringField, SelectField, DateTimeField, PasswordField, validators
from passlib.hash import sha256_crypt
from functools import wraps
from datetime import datetime as dt
from datetime import timedelta
from werkzeug.utils import secure_filename
import pytz
import random
import string
from collections import Counter
import pandas as pd
import re
from docx import Document
import os

from io import BytesIO
import io
from googleapiclient.http import MediaIoBaseUpload
from googleapiclient.http import MediaIoBaseDownload
from googleapiclient.errors import HttpError
import requests

from classes import *
from globals import *
from gdrive import GoogleDriveService

import h5py
import numpy as np
from tensorflow.keras.models import load_model
from tensorflow.keras.optimizers import Adam

app.jinja_env.tests['equalto'] = lambda value, other : value == other

# Aliases for readability in the join
UserAlias = aliased(User)
DoctorAlias = aliased(Doctor)

# Connect to the database
engine = create_engine(database_url)

# Test the connection
connection = engine.connect()

# Create database tables
with app.app_context():
    db.create_all()

# Set parameters manually
path_to_model = "model_files/model.hdf5"
dataset_name = "tracings"
batch_size = 32

#Email parameters
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 465
app.config['MAIL_USE_TLS'] = False
app.config['MAIL_USE_SSL'] = True
app.config['MAIL_USERNAME'] = 'andriylv94@gmail.com'
app.config['MAIL_PASSWORD'] = 'gdbk efjw clpp edgf'

mail = Mail(app)

ALLOWED_EXTENSIONS = {'docx'}

@app.route('/')
def index():
    return render_template('home.html')


# Register class
class RegisterForm(Form):
    name = StringField('Ім\'я', [validators.Length(min=1, max=40)])
    surname = StringField('Прізвище', [validators.Length(min=4, max=50)])
    patronymic = StringField('По батькові', [validators.Length(min=4, max=40)])
    gender_choices = [('MALE', 'чоловік'), ('FEMALE', 'жінка')]  # Enum of male or female
    gender = SelectField('Стать', choices=gender_choices)
    email = StringField('Електронна пошта', [validators.Length(min=6, max=255)])
    phone_number = StringField('Номер телефону', [validators.Length(min=6, max=20)])
    password = PasswordField('Пароль', [
        validators.DataRequired(),
        validators.EqualTo('confirm', message="Password do not match"),
    ])
    confirm = PasswordField("Підтвердити пароль")


def generate_verification_code():
    return ''.join(random.choice(string.ascii_uppercase + string.digits) for _ in range(6))


# Register User
@app.route('/register', methods=['GET', 'POST'])
def register():
    form = RegisterForm(request.form)
    if request.method == 'POST' and form.validate():
        name = form.name.data
        surname = form.surname.data
        patronymic = form.patronymic.data
        gender = form.gender.data
        email = form.email.data
        password = sha256_crypt.encrypt(form.password.data)
        phone_number = form.phone_number.data

        if User.query.filter_by(email=email).first():
            flash('This email is already taken', 'danger')
            return render_template('register.html', form=form)

        session['temp_user'] = {
            'name': name,
            'surname': surname,
            'patronymic': patronymic,
            'gender': gender,
            'email': email,
            'password': password,  # Encrypted password
            'phone_number': phone_number
        }

        # Generate and send verification code
        verification_code = generate_verification_code()
        msg = Message('Your Verification Code', sender='your-email@example.com', recipients=[email])
        msg.body = f'Your verification code is {verification_code}'
        mail.send(msg)

        # Store verification code in the session or database
        session['verification_code'] = verification_code  # Example: storing in session

        # Render a template that asks for the verification code
        return render_template('verify.html', email=email, form=form)

    return render_template('register.html', form=form)


@app.route('/verify', methods=['POST'])
def verify():
    entered_code = request.form['verification_code']

    if 'verification_code' in session and session['verification_code'] == entered_code:
        # Retrieve the stored user data
        temp_user = session.pop('temp_user', None)

        if temp_user:
            # Proceed with creating the User object and saving it to the database
            user = User(**temp_user)
            db.session.add(user)
            db.session.commit()

            flash('You are now registered and can log in', 'success')
            return redirect(url_for('login'))
        else:
            flash('Registration session expired or invalid', 'danger')
            return redirect(url_for('register'))
    else:
        flash('Invalid verification code', 'danger')
        return redirect(url_for('register'))


@app.route('/login', methods=["GET", "POST"])
def login():
    if request.method == 'POST':
        # Get Form Fields
        email = request.form['email']
        password_candidate = request.form['password']

        # Check if email exists in 'user' table
        user = User.query.filter_by(email=email).first()
        if user:
            role = 'user'
        else:
            # Check if email exists in 'doctor' table
            doctor = Doctor.query.filter_by(email=email).first()
            if doctor:
                role = 'doctor'
            else:
                # Check if email exists in 'admin' table
                admin = Admin.query.filter_by(email=email).first()
                if admin:
                    role = 'admin'
                else:
                    # Email not found in any table
                    error = 'Уведена електронна пошта не знайдена'
                    return render_template('login.html', error=error)

        # Based on the role, check password and set session role
        if role == 'user':
            user_obj = user
        elif role == 'doctor':
            user_obj = doctor
        else:
            user_obj = admin

        # Get stored hash
        password_hash = user_obj.password

        # Compare Password
        if sha256_crypt.verify(password_candidate, password_hash):
            # Passed
            session['logged_in'] = True
            session['user_id'] = user_obj.id
            session['role'] = role

            flash('Ви увійшли', 'success')
            return redirect(url_for('index'))
        else:
            error = 'Неправильний пароль'
            return render_template('login.html', error=error)

    return render_template('login.html')


# Check if user is logged in
def is_logged_in(f):
    @wraps(f)
    def wrap(*args, **kwargs):
        if 'logged_in' in session:
            return f(*args, **kwargs)
        else:
            flash('Ви не увійшли, будь ласка, увійдіть', 'danger')
            return redirect(url_for('login'))

    return wrap


# Check if user is user
def is_user(f):
    @wraps(f)
    def wrap(*args, **kwargs):
        if session['role'] == 'user':
            return f(*args, **kwargs)
        else:
            flash('Ви не маєте доступу до цієї сторінки!', 'danger')
            return redirect(url_for('index'))

    return wrap


# Check if user is doctor
def is_doctor(f):
    @wraps(f)
    def wrap(*args, **kwargs):
        if session['role'] == 'doctor':
            return f(*args, **kwargs)
        else:
            flash('Ви не маєте доступу до цієї сторінки', 'danger')
            return redirect(url_for('index'))

    return wrap


# Check if user is admin
def is_admin(f):
    @wraps(f)
    def wrap(*args, **kwargs):
        if session['role'] == 'admin':
            return f(*args, **kwargs)
        else:
            flash('Ви не маєте доступу до цієї сторінки', 'danger')
            return redirect(url_for('index'))

    return wrap


def is_doctor_or_user(f):
    @wraps(f)
    def wrap(*args, **kwargs):
        if session['role'] == 'doctor' or session['role'] == 'user':
            return f(*args, **kwargs)
        else:
            flash('Ви не маєте доступу до цієї сторінки', 'danger')
            return redirect(url_for('index'))

    return wrap


# Logout
@app.route('/logout')
@is_logged_in
def logout():
    session.clear()
    flash("Ви вийшли", 'success')
    return redirect(url_for('login'))


@app.route('/dashboard')
@is_logged_in
@is_doctor_or_user
def dashboard():
    if session['role'] == 'doctor':
        doctor = Doctor.query.filter_by(id=session['user_id']).first()
        appointments = Appointment.query.filter_by(doctor_id=session['user_id']).all()
        users = User.query.all()
        hospitals = Hospital.query.all()
        if appointments:
            return render_template("dashboard_doctor.html", users=users, doctor=doctor,
                                   appointments=appointments, hospitals=hospitals)
        else:
            msg = 'До вас ще не записувались пацієнти'
            return render_template("dashboard_doctor.html", msg=msg)
    elif session['role'] == 'user':
        # Query appointments for the current user
        appointments = Appointment.query.filter_by(user_id=session['user_id']).all()
        hospitals = Hospital.query.all()
        doctors = Doctor.query.all()
        if appointments:
            return render_template("dashboard.html", appointments=appointments, hospitals=hospitals,
                                   doctors=doctors)
        else:
            msg = 'У вас ще не було записів до лікарів'
            return render_template("dashboard.html", msg=msg)


@app.route('/delete_appointment/<id>', methods=['POST'])
@is_logged_in
def delete_appointment(id):
    appointment = Appointment.query.get(id)
    if appointment:
        db.session.delete(appointment)
        db.session.commit()
    return redirect(url_for('dashboard'))


@app.route('/confirm_delete_appointment/<id>', methods=['POST'])
@is_logged_in
def confirm_delete_appointment(id):
    if request.method == 'POST':
        user_obj = User.query.filter_by(id=session['user_id']).first()
        # Get stored hash
        password_hash = user_obj.password
        # Verify password here
        password = request.form.get('password')
        if sha256_crypt.verify(password, password_hash):
            appointment = Appointment.query.get(id)
            if appointment:
                db.session.delete(appointment)
                db.session.commit()
            return redirect(url_for('dashboard'))
        else:
            flash('Неправильний пароль. Видалення скасовано.', 'danger')
            return redirect(url_for('dashboard'))


@app.route('/decline_appointment/<int:appointment_id>', methods=['GET'])
@is_logged_in
def decline_appointment(appointment_id):
    if session['role'] == 'doctor':
        appointment = Appointment.query.get(appointment_id)
        if appointment:
            appointment.status = 'Відхилено'
            appointment.date_time = None
            db.session.commit()
            flash('Запис відхилено', 'success')
        else:
            flash('Запис не знайдено', 'danger')
        return redirect(url_for('appointment', id=appointment_id))
    else:
        flash('Немає доступу', 'danger')
        return redirect(url_for('dashboard'))


@app.route('/users')
@is_logged_in
@is_admin
def users():
    if session['role'] == 'admin':
        users_list = User.query.all()
        doctors = Doctor.query.all()
        admins = Admin.query.all()
        hospitals_list = Hospital.query.all()
        return render_template("users.html", users=users_list, doctors=doctors, admins=admins, hospitals=hospitals_list)
    else:
        return redirect(url_for('dashboard'))


@app.route('/confirm_delete_user/<id>', methods=['POST'])
@is_logged_in
@is_admin
def confirm_delete_user(id):
    if request.method == 'POST':
        admin_obj = Admin.query.filter_by(id=session['user_id']).first()
        # Get stored hash
        password_hash = admin_obj.password
        # Verify password here
        password = request.form.get('password')
        if sha256_crypt.verify(password, password_hash):
            user = User.query.get(id)
            if user:
                db.session.delete(user)
                db.session.commit()
            return redirect(url_for('users'))
        else:
            flash('Неправильний пароль. Видалення скасовано.')
            return redirect(url_for('users'))


@app.route('/confirm_delete_doctor/<id>', methods=['POST'])
@is_logged_in
@is_admin
def confirm_delete_doctor(id):
    if request.method == 'POST':
        admin_obj = Admin.query.filter_by(id=session['user_id']).first()
        # Get stored hash
        password_hash = admin_obj.password
        # Verify password here
        password = request.form.get('password')
        if sha256_crypt.verify(password, password_hash):
            doctor = Doctor.query.get(id)
            if doctor:
                db.session.delete(doctor)
                db.session.commit()
            return redirect(url_for('users'))
        else:
            flash('Неправильний пароль. Видалення скасовано.')
            return redirect(url_for('users'))


@app.route('/promote_to_doctor/<int:user_id>/<int:hospital_id>', methods=['POST'])
@is_logged_in
@is_admin
def promote_to_doctor(user_id, hospital_id):
    user = User.query.get(user_id)
    hospital = Hospital.query.get(hospital_id)
    if user and hospital:
        doctor = Doctor(
            name=user.name,
            surname=user.surname,
            patronymic=user.patronymic,
            gender=user.gender,
            email=user.email,
            phone_number=user.phone_number,
            password=user.password,  # Ensure this is already hashed
            hospital=hospital
        )
        db.session.add(doctor)
        db.session.delete(user)
        db.session.commit()
        flash('User has been promoted to doctor', 'success')
    else:
        flash('User or hospital not found', 'danger')
    return redirect(url_for('users'))



@app.route('/select_hospital_for_doctor/<int:user_id>', methods=['GET'])
@is_logged_in
@is_admin
def select_hospital_for_doctor(user_id):
    hospitals = Hospital.query.all()
    return render_template('select_hospital_for_doctor.html', hospitals=hospitals, user_id=user_id)


@app.route('/demote_to_user/<id>', methods=['POST'])
@is_logged_in
@is_admin
def demote_to_user(id):
    if request.method == 'POST':
        admin_obj = Admin.query.filter_by(id=session['user_id']).first()
        password_hash = admin_obj.password
        password = request.form.get('password')
        if sha256_crypt.verify(password, password_hash):
            doctor = Doctor.query.get(id)
            if doctor:
                # Create a new User object using the doctor's data
                user = User(
                    name=doctor.name,
                    surname=doctor.surname,
                    patronymic=doctor.patronymic,
                    gender=doctor.gender,
                    email=doctor.email,
                    phone_number=doctor.phone_number,
                    password=doctor.password
                )
                db.session.add(user)
                db.session.delete(doctor)  # Remove the doctor
                db.session.commit()
            return redirect(url_for('users'))
        else:
            flash('Неправильний пароль. Операцію відмінено.')
            return redirect(url_for('users'))


@app.route('/hospitals', methods=['GET'])
@is_logged_in
@is_admin
def hospitals():
    search_query = request.args.get('search')
    if search_query:
        hospitals = Hospital.query.filter(Hospital.name.contains(search_query)).all()
        if hospitals:
            return render_template('hospitals.html', hospitals=hospitals)
        else:
            return render_template('hospitals.html', msg='Лікарень не знайдено')
    else:
        hospitals = Hospital.query.all()
        return render_template('hospitals.html', hospitals=hospitals)


class AddHospitalForm(Form):
    name = StringField('Назва', [validators.Length(min=1, max=150)])
    location = StringField('Адреса', [
        validators.Regexp(r'^вулиця [^\,]+, \d+, [^\,]+$',
                          message='Адреса повинна бути у форматі: "вулиця *назва вулиці*, *номер будинку*, *місто*"')
    ])
    contact_number = StringField('Номер телефону', [
        validators.Regexp(r'^\+38\(\d{3}\)\d{3}-\d{2}-\d{2}$',
                          message='Номер телефону повинен мати формат +38(XXX)XXX-XX-XX')
    ])


@app.route('/add_hospital', methods=['GET', 'POST'])
@is_logged_in
@is_admin
def add_hospital():
    form = AddHospitalForm(request.form)
    if request.method == 'POST' and form.validate():
        name = form.name.data
        location = form.location.data
        contact_number = form.contact_number.data
        new_hospital = Hospital(name=name, location=location, contact_number=contact_number)
        db.session.add(new_hospital)
        db.session.commit()
        flash('Лікарня додана', 'success')
        return redirect(url_for('hospitals'))
    return render_template('add_hospital.html', form=form)


@app.route('/delete_hospital/<int:id>', methods=['GET', 'POST'])
@is_logged_in
@is_admin
def delete_hospital(id):
    hospital = Hospital.query.get_or_404(id)
    if hospital:
        # Process each doctor related to the hospital
        doctors = Doctor.query.filter_by(hospital_id=hospital.id).all()
        for doctor in doctors:
            appointments = Appointment.query.filter_by(doctor_id=doctor.id).all()
            for appointment in appointments:
                if appointment.status != 'Перевірено':
                    appointment.status = 'Відхилено'
                appointment.location = None
                appointment.doctor_id = None
                db.session.commit()

            ecgs = ECG.query.filter_by(doctor_id=doctor.id).all()
            for ecg in ecgs:
                # Delete independent ECGs not linked to any user
                if ecg.user_id is None:
                    if ecg.ecg_file:
                        file_id = ecg.ecg_file.split('/')[-2]  # Assuming ecg.ecg_file stores the file ID
                        if delete_file_from_google_drive(file_id):
                            ecg.ecg_file = None
                    db.session.delete(ecg)
                else:
                    ecg.doctor_id = None
                    db.session.commit()

            # Convert doctors to users if necessary or delete them
            new_user = User(name=doctor.name, surname=doctor.surname, patronymic=doctor.patronymic,
                            gender=doctor.gender, email=doctor.email, password=doctor.password,
                            phone_number=doctor.phone_number)
            db.session.add(new_user)
            db.session.delete(doctor)
            db.session.commit()

        # Delete the hospital
        db.session.delete(hospital)
        db.session.commit()
        flash('Hospital and all related data safely updated.', 'success')
    else:
        flash('Hospital not found.', 'danger')

    return redirect(url_for('hospitals'))


def delete_file_from_google_drive(file_id):
    try:
        google_drive_service = GoogleDriveService().build()
        google_drive_service.files().delete(fileId=file_id).execute()
        return True
    except Exception as e:
        print(f"Failed to delete file: {e}")
        return False


@app.route('/confirm_delete_hospital/<id>', methods=['POST'])
@is_logged_in
@is_admin
def confirm_delete_hospital(id):
    if request.method == 'POST':
        admin_obj = Admin.query.filter_by(id=session['user_id']).first()
        # Get stored hash
        password_hash = admin_obj.password
        # Verify password here
        password = request.form.get('password')
        if sha256_crypt.verify(password, password_hash):
            return redirect(url_for('delete_hospital', id=id))
        else:
            flash('Неправильний пароль. Видалення скасовано.')
            return redirect(url_for('hospitals'))


# profile page
@app.route('/profile')
@is_logged_in
def profile():
    if session['role'] == 'user':
        user = User.query.filter_by(id=session['user_id']).first()
        return render_template('profile.html', user=user, Gender=Gender)
    elif session['role'] == 'doctor':
        doctor = Doctor.query.filter_by(id=session['user_id']).first()
        hospital = Hospital.query.filter_by(id=doctor.hospital_id).first()
        hospital_name = hospital.name if hospital else None
        return render_template('profile.html', user=doctor, Gender=Gender, hospital_name=hospital_name)
    else:
        admin = Admin.query.filter_by(id=session['user_id']).first()
        return render_template('profile.html', user=admin, Gender=Gender)


@app.route('/update_profile', methods=['POST'])
@is_logged_in
def update_profile():
    user_id = session['user_id']
    user_role = session['role']

    if user_role == 'user':
        user = User.query.filter_by(id=user_id).first()
    elif user_role == 'doctor':
        user = Doctor.query.filter_by(id=user_id).first()
    else:
        user = Admin.query.filter_by(id=user_id).first()

    if user:
        # Get the data from the form
        user.name = request.form['name']
        user.surname = request.form['surname']
        user.patronymic = request.form['patronymic']
        user.gender = request.form['gender']
        user.phone_number = request.form['phone_number']
        new_email = request.form['email']

        # Check if the new email is not taken yet
        if user.email != new_email and (User.query.filter_by(email=new_email).first() or Doctor.query.filter_by(email=new_email).first() or Admin.query.filter_by(email=new_email).first()):
            flash('Email is already taken. Please choose a different one.', 'danger')
            return redirect(url_for('profile'))

        user.email = new_email
        db.session.commit()
        flash('Профіль успішно оновлений', 'success')
    else:
        flash('Користувач не знайдений', 'danger')

    return redirect(url_for('profile'))


@app.route('/change_password', methods=['POST'])
@is_logged_in
def change_password():
    user_id = session['user_id']
    user_role = session['role']

    if user_role == 'user':
        user = User.query.filter_by(id=user_id).first()
    elif user_role == 'doctor':
        user = Doctor.query.filter_by(id=user_id).first()
    else:
        user = Admin.query.filter_by(id=user_id).first()

    if user:
        current_password = request.form['current_password']
        new_password = request.form['new_password']
        confirm_password = request.form['confirm_password']

        # Check if current password matches
        if not sha256_crypt.verify(current_password, user.password):
            flash('Неправильний поточний пароль.', 'danger')
            return redirect(url_for('profile'))

        # Check if new password and confirm password match
        if new_password != confirm_password:
            flash('Новий пароль і підтвердження пароля не співпадають.', 'danger')
            return redirect(url_for('profile'))

        # Update the password
        user.password = sha256_crypt.encrypt(new_password)
        db.session.commit()
        flash('Пароль змінено успішно.', 'success')
    else:
        flash('Користувач не знайдений.', 'danger')

    return redirect(url_for('profile'))


@app.route('/add_appointment', methods=['GET', 'POST'])
@is_logged_in
@is_user
def add_appointment():
    search_query = request.args.get('search')
    if search_query:
        hospitals = Hospital.query.filter(Hospital.name.contains(search_query)).all()
    else:
        hospitals = Hospital.query.all()
    return render_template('add_appointment.html', hospitals=hospitals)


@app.route('/choose_hospital/<int:hospital_id>', methods=['GET'])
@is_logged_in
@is_user
def choose_hospital(hospital_id):
    hospital = Hospital.query.get(hospital_id)
    if hospital:
        doctors = Doctor.query.filter_by(hospital_id=hospital_id).all()
        return render_template('choose_doctor.html', hospital=hospital, doctors=doctors)
    else:
        return redirect(url_for('add_appointment', msg='Лікарню не знайдено'))


@app.route('/book_appointment/<int:doctor_id>', methods=['POST'])
@is_logged_in
@is_user
def book_appointment(doctor_id):
    doctor = Doctor.query.get(doctor_id)
    if doctor:
        appointment = Appointment(
            location=doctor.hospital.location,
            date_time=None,  # To be set by the doctor
            status='В очікуванні',
            doctor_id=doctor_id,
            user_id=session['user_id'],
            ecg_id=None
        )

        db.session.add(appointment)
        db.session.commit()
        return redirect(url_for('dashboard'))
    else:
        return redirect(url_for('choose_hospital'))


@app.route('/appointment/<id>')
@is_logged_in
def appointment(id):
    appointment = Appointment.query.filter_by(id=id).first()
    if appointment:
        if not appointment.doctor_id:
            doctor = None
        else:
            doctor = Doctor.query.filter_by(id=appointment.doctor_id).first()
        if doctor == None:
            hospital = None
        else:
            hospital = Hospital.query.filter_by(id=doctor.hospital_id).first()
        user = User.query.filter_by(id=appointment.user_id).first()
        ecg = ECG.query.filter_by(id=appointment.ecg_id).first()
        return render_template('appointment.html', appointment=appointment, user=user, doctor=doctor, hospital=hospital, ecg=ecg)
    else:
        return redirect(url_for('dashboard'))


@app.route('/set_appointment_time/<int:appointment_id>', methods=['POST'])
@is_logged_in
def set_appointment_time(appointment_id):
    if session['role'] == 'doctor':
        appointment = Appointment.query.get(appointment_id)
        if appointment:
            appointment_time_str = request.form['appointment_time']
            appointment_date = dt.strptime(appointment_time_str, '%Y-%m-%dT%H:%M').date()
            appointment_time = dt.strptime(appointment_time_str, '%Y-%m-%dT%H:%M').time()
            appointment_date_time = dt.combine(appointment_date, appointment_time)

            # Get current time in Ukraine
            tz = pytz.timezone('Europe/Kiev')
            current_date = dt.now(tz).date()
            current_time = dt.now(tz).time()
            current_date_time = dt.combine(current_date, current_time)

            if appointment_date_time <= current_date_time:
                flash('Дата і час мають бути більші за поточний момент', 'danger')
                return redirect(url_for('appointment', id=appointment_id))

            appointment.date_time = appointment_date_time
            appointment.status = 'Призначено'
            db.session.commit()
            flash('Дата і час призначені', 'success')
            return redirect(url_for('appointment', id=appointment_id))
        else:
            flash('Запис не знайдено', 'danger')
            return redirect(url_for('dashboard'))
    else:
        flash('Немає доступу', 'danger')
        return redirect(url_for('dashboard'))


# route to upload ecg or choose the user to upload ecg, doctor can upload ecg even without selecting a user
@app.route('/upload_ecg', methods=['GET', 'POST'])
@is_logged_in
@is_doctor
def upload_ecg():
    if request.method == 'POST':
        user_id = request.form.get('user_id')
        user = User.query.get(user_id)
        doctor = Doctor.query.get(session['user_id'])
        ecg_file = request.files['ecg_file']
        # store appointemnt id
        appointment_id = request.form.get('appointment_id')

        if ecg_file:
            # save ecg_file to the google drive and add the bew link to the ecg_file to database
            # start code here
            buffer_memory = BytesIO()
            ecg_file.save(buffer_memory)

            # Create Google Drive service object
            google_drive_service = GoogleDriveService().build()

            # File metadata for Google Drive
            file_metadata = {
                'name': ecg_file.filename,
                'parents': ['1NiKjW8pGf1TPBau2cgV_pp5xcMEGVwbO']  # Folder ID
            }

            # Media content for Google Drive
            media = MediaIoBaseUpload(ecg_file, mimetype=ecg_file.mimetype, resumable=True)

            returned_fields = "id, name, mimeType, webViewLink, exportLinks"
            # Upload file to Google Drive
            file = google_drive_service.files().create(body=file_metadata, media_body=media,
                                                       fields=returned_fields).execute()

            # Get the webViewLink to store in the database
            web_view_link = file.get('webViewLink')

            # Create a permission to make the file viewable by anyone with the link
            permission = {
                'type': 'anyone',
                'role': 'reader',
            }

            try:
                google_drive_service.permissions().create(fileId=file['id'], body=permission).execute()
            except HttpError as error:
                print(f'An error occurred: {error}')

            # Store the ECG information including the webViewLink in the database
            new_ecg = ECG(
                ecg_file=web_view_link,
                datetime=dt.now(),
                results=None,
                doctor=doctor,
                user=user
            )
            db.session.add(new_ecg)
            # change the ecg_id in the corresponding appointment to the new ecg_id
            if new_ecg.doctor and new_ecg.user:
                appointment = Appointment.query.get(appointment_id)
                appointment.ecg_id = new_ecg.id
            db.session.commit()
            flash('ECG uploaded successfully', 'success')
            return redirect(url_for('dashboard'))

        else:
            flash('No file uploaded', 'danger')
            return redirect(url_for('upload_ecg'))
    else:
        users = User.query.all()
        appointments = Appointment.query.filter_by(doctor_id=session['user_id']).all()
        return render_template('upload_ecg.html', users=users, appointments=appointments)


def delete_file_from_google_drive(file_id):
    try:
        google_drive_service = GoogleDriveService().build()
        google_drive_service.files().delete(fileId=file_id).execute()
        return True
    except Exception as e:
        print(f"Failed to delete file: {e}")
        return False


# check_ecg route where the table with all ecgs is displayed
@app.route('/check_ecg')
@is_logged_in
@is_doctor
def check_ecg():
    # Subquery to find all ecg_ID values in the appointment table
    subquery = db.session.query(Appointment.ecg_id).filter(Appointment.ecg_id == ECG.id).subquery()
    # Query to find all ECG records that are NOT IN the subquery results
    ecgs = ECG.query.filter(
        not_(ECG.id.in_(subquery)),
        ECG.doctor_id == session['user_id']
    ).order_by(ECG.datetime.desc()).all()
    appointments = Appointment.query.filter_by(doctor_id=session['user_id']).order_by(Appointment.date_time.desc()).all()
    return render_template('check_ecg.html', ecgs=ecgs, appointments=appointments)


@app.route('/ecg/<id>')
@is_logged_in
@is_doctor
def ecg(id):
    ecg = ECG.query.filter_by(id=id).first()
    if ecg:
        return render_template('ecg.html', ecg=ecg)
    else:
        flash('ECG not found.', 'danger')
        return redirect(url_for('check_ecg'))


@app.route('/get_ecg_data/<id>/<segment>')
@is_logged_in
@is_doctor
def get_ecg_data(id, segment):
    ecg = ECG.query.filter_by(id=id).first()
    if ecg:
        # Create Google Drive service object
        google_drive_service = GoogleDriveService().build()

        # Replace 'webViewLink' with the 'id' of the file if you store only the ID in the database
        file_id = ecg.ecg_file.split('/')[-2]

        request = google_drive_service.files().get_media(fileId=file_id)
        buffer_memory = BytesIO()
        downloader = MediaIoBaseDownload(buffer_memory, request)

        done = False
        while done is False:
            status, done = downloader.next_chunk()

        buffer_memory.seek(0)

        # Load the HDF5 file
        loaded_file = h5py.File(buffer_memory, 'r')
        ecg_data = np.array(loaded_file['tracings'])

        # Extracting the specified segment for all leads
        segment_index = int(segment) - 1  # Adjust for zero-indexing if necessary
        selected_segment_all_leads = ecg_data[segment_index, :, :].T  # Transpose to list leads separately
        selected_segment_all_leads = selected_segment_all_leads.tolist()  # Convert to list of lists

        # Construct a response with time points and voltage values for the selected segment
        ecg_plot_data = {
            'x_values': list(range(ecg_data.shape[1])),  # Assuming second dimension is time
            'y_values': selected_segment_all_leads
        }
        return jsonify(ecg_plot_data)
    else:
        return abort(404)  # Not found


@app.route('/save_ecg_result/<id>', methods=['POST'])
@is_logged_in
@is_doctor
def save_ecg_result(id):
    ecg = ECG.query.filter_by(id=id).first()
    if ecg:
        selected_results = request.form.getlist('results')  # Gets all checked options
        ecg.results = ', '.join(selected_results)  # Join the results into a single string
        # if this ecg has connected appointment, change corresponding appointment record status to 'Перевірено'
        appointment = Appointment.query.filter_by(ecg_id=id).first()
        if appointment:
            appointment.status = 'Перевірено'
        db.session.commit()
        flash('ECG results updated successfully.', 'success')
    else:
        flash('ECG not found.', 'danger')

    return redirect(url_for('ecg', id=id))


@app.route('/auto_check/<id>')
@is_logged_in
@is_doctor
def auto_check(id):
    ecg = ECG.query.filter_by(id=id).first()
    if ecg:
        # Create Google Drive service object
        google_drive_service = GoogleDriveService().build()

        # Replace 'webViewLink' with the 'id' of the file if you store only the ID in the database
        file_id = ecg.ecg_file.split('/')[-2]

        request = google_drive_service.files().get_media(fileId=file_id)
        buffer_memory = BytesIO()
        downloader = MediaIoBaseDownload(buffer_memory, request)

        done = False
        while done is False:
            status, done = downloader.next_chunk()

        buffer_memory.seek(0)

        # Load the HDF5 file
        loaded_file = h5py.File(buffer_memory, 'r')
        ecg_data = np.array(loaded_file['tracings'])
        model = load_model(path_to_model, compile=False)
        model.compile(loss='binary_crossentropy', optimizer=Adam())
        y_score = model.predict(ecg_data, verbose=1)
        return jsonify(y_score.tolist())
    else:
        return jsonify([]), 404


@app.route('/statistics')
@is_logged_in
@is_admin
def statistics():
    # Get all ECG records and split the results by comma to count each outcome
    ecgs = ECG.query.all()
    outcomes = Counter()
    for ecg in ecgs:
        if ecg.results:
            for outcome in ecg.results.split(', '):
                outcomes[outcome.strip()] += 1

    # Get the number of users that had appointments in each hospital
    hospital_appointments = db.session.query(Hospital.id, Hospital.name, db.func.count(Appointment.id).label('appointments_count')) \
                                       .join(Doctor, Hospital.id == Doctor.hospital_id) \
                                       .join(Appointment, Doctor.id == Appointment.doctor_id) \
                                       .group_by(Hospital.id) \
                                       .all()

    return render_template('statistics.html', outcomes=outcomes, hospital_appointments=hospital_appointments)


@app.route('/export_hospitals')
@is_logged_in
@is_admin
def export_hospitals():
    # Query hospital data from the database
    hospitals = Hospital.query.all()
    # Convert the hospital data to a pandas DataFrame
    data = {
        "Name": [hospital.name for hospital in hospitals],
        "Location": [hospital.location for hospital in hospitals],
        "ContactNumber": [hospital.contact_number for hospital in hospitals]
    }
    df = pd.DataFrame(data)

    # Save the DataFrame to an Excel file in memory
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Hospitals')

    # Rewind the buffer
    output.seek(0)

    # Send the file for download
    return send_file(output, as_attachment=True, download_name='hospitals.xlsx',
                     mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in {'xlsx'}


@app.route('/import_hospitals', methods=['POST'])
@is_logged_in
@is_admin
def import_hospitals():
    if 'file' not in request.files:
        flash('No file part', 'danger')
        return redirect(url_for('hospitals'))
    file = request.files['file']
    if file.filename == '':
        flash('No selected file', 'danger')
        return redirect(url_for('hospitals'))
    if file and not allowed_file(file.filename):
        try:
            df = pd.read_excel(file)
            errors = validate_and_import(df)
            if errors:
                flash(f'Errors occurred: {", ".join(errors)}', 'danger')
            else:
                flash('Лікарні успішно завантажені', 'success')
        except Exception as e:
            flash(f'An error occurred while reading the file: {str(e)}', 'danger')
    else:
        flash('Invalid file type, please upload an Excel file.', 'danger')
    return redirect(url_for('hospitals'))


def validate_and_import(df):
    errors = []
    for index, row in df.iterrows():
        name = row['Name']
        location = row['Location']
        contact_number = row['ContactNumber']

        # Validate phone number
        if not re.match(r'^\+38\(\d{3}\)\d{3}-\d{2}-\d{2}$', contact_number):
            errors.append(f'Invalid contact number format at row {index + 1}')
            continue

        # Validate location
        if not re.match(r'^вулиця [^\,]+, \d+, [^\,]+$', location):
            errors.append(f'Invalid location format at row {index + 1}')
            continue

        # Check if hospital already exists
        if not Hospital.query.filter_by(name=name, location=location, contact_number=contact_number).first():
            new_hospital = Hospital(name=name, location=location, contact_number=contact_number)
            db.session.add(new_hospital)

    if not errors:
        db.session.commit()

    return errors


@app.route('/edit_hospital/<int:id>', methods=['GET', 'POST'])
@is_logged_in
@is_admin
def edit_hospital(id):
    hospital = Hospital.query.get_or_404(id)
    form = AddHospitalForm(request.form)

    if request.method == 'POST' and form.validate():
        hospital.name = form.name.data
        hospital.location = form.location.data
        hospital.contact_number = form.contact_number.data

        db.session.commit()
        flash('Лікарню оновлено', 'success')
        return redirect(url_for('hospitals'))

    form.name.data = hospital.name
    form.location.data = hospital.location
    form.contact_number.data = hospital.contact_number

    return render_template('edit_hospital.html', form=form, hospital_id=hospital.id)


@app.route('/generate_ecg_report')
@is_logged_in
@is_admin
def generate_ecg_report():
    # Calculate date range for the last 30 days
    end_date = dt.today()  # Current date
    start_date = end_date - timedelta(days=30)  # Date 30 days ago

    # Query the database for ECG records within the last 30 days
    ecgs = ECG.query \
        .join(Doctor, ECG.doctor_id == Doctor.id) \
        .join(User, ECG.user_id == User.id) \
        .add_columns(User.name.label('patient_name'), User.surname.label('patient_surname'),
                     User.patronymic.label('patient_patronymic'), ECG.datetime, Doctor.name.label('doctor_name'),
                     Doctor.surname.label('doctor_surname'), Doctor.patronymic.label('doctor_patronymic'), ECG.results) \
        .filter(ECG.datetime >= start_date, ECG.datetime <= end_date).all()

    # Create a new Document
    doc = Document()
    doc.add_heading('ЕКГ звіт за останні 30 днів', level=1)
    doc.add_paragraph(f'Звіт за період від: {start_date.strftime("%Y-%m-%d")} до {end_date.strftime("%Y-%m-%d")}')
    # Add a table to the document
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ПІБ пацієнта'
    hdr_cells[1].text = 'Дата ЕКГ'
    hdr_cells[2].text = 'ПІБ лікаря'
    hdr_cells[3].text = 'Результат'

    # Fill table with ECG data
    for ecg in ecgs:
        row_cells = table.add_row().cells
        # Access the results directly since they are not attributes of an object
        patient_full_name = f"{ecg.patient_surname} {ecg.patient_name} {ecg.patient_patronymic}"
        doctor_full_name = f"{ecg.doctor_surname} {ecg.doctor_name} {ecg.doctor_patronymic}"

        row_cells[0].text = patient_full_name
        row_cells[1].text = ecg.datetime.strftime("%Y-%m-%d / %H:%M:%S")
        row_cells[2].text = doctor_full_name
        row_cells[3].text = ecg.results

    # Save the document
    file_path = "ECG_Report.docx"
    doc.save(file_path)

    # Send file to user
    return send_file(file_path, as_attachment=True, download_name='Monthly_ECG_Report.docx')


def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


@app.route('/upload_ecg_report', methods=['GET','POST'])
@is_logged_in
@is_admin
def upload_ecg_report():
    # Check if 'report' is in the files part of the request
    if 'report' not in request.files:
        flash('No file part')
        return redirect(request.url)
    file = request.files['report']
    # If the user does not select a file, the browser submits an
    # empty file without a filename.
    if file.filename == '':
        flash('No selected file')
        return redirect(request.url)
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join('', filename)
        file.save(file_path)
        try:
            doc = Document(file_path)
            updated_records_count = 0
            for table in doc.tables:
                # Assume we are iterating over each row in the .docx table
                for row in table.rows[1:]:  # Skip header row
                    patient_full_name = row.cells[0].text.strip()
                    date_of_ecg = row.cells[1].text.strip()
                    doctor_full_name = row.cells[2].text.strip()
                    outcome = row.cells[3].text.strip()
                    # Parse the date string into a datetime object
                    ecg_date = dt.strptime(date_of_ecg, "%Y-%m-%d / %H:%M:%S")
                    # Find the ECG record using the parsed name parts and date
                    ecg_record = ECG.query.where(ECG.datetime == ecg_date).first()
                    # Update the record if it exists
                    if ecg_record:
                        ecg_record.results = outcome
                        db.session.commit()
                    else:
                        continue
                    updated_records_count += 1  # Increment the count of updated records
            if updated_records_count == 0:
                flash('There are no records that can be changed using the report', 'warning')
            else:
                flash(f'{updated_records_count} records have been updated from the report', 'success')
        except Exception as e:
            flash(str(e), 'danger')
        # Finally, remove the uploaded file to clean up
        os.remove(file_path)
        return redirect(url_for('statistics'))  # Replace 'statistics' with the correct route for your statistics page
    else:
        flash('Invalid file type', 'danger')
        return redirect(request.url)


if __name__ == '__main__':
    app.secret_key = 'secret123'
    app.run(host='0.0.0.0', debug=True)